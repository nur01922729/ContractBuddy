import fitz  # PyMuPDF
# import easyocr
import io
from PIL import Image
from groq import Groq
import os
from dotenv import load_dotenv
import json

load_dotenv()
client = Groq(api_key=os.getenv("GROQ_API_KEY"))

# def extract_text_from_file(file_path: str) -> str:
#     print(f"Extracting from: {file_path}")
    
#     text = ""
#     ext = file_path.lower().split(".")[-1]

#     try:
#         # Handle temporary files (.tmp) by treating them as their original type or using PyMuPDF
#         if ext in ["tmp", "temp"]:
#             # Try PyMuPDF first (works for most PDFs and some docs)
#             try:
#                 doc = fitz.open(file_path)
#                 for page in doc:
#                     text += page.get_text("text") + "\n"
#                 doc.close()
#                 print("PyMuPDF extraction successful for temp file")
#             except:
#                 # Fallback to OCR if PyMuPDF fails
#                 print("PyMuPDF failed, trying OCR...")
#                 reader = easyocr.Reader(['en', 'hi'], gpu=False)
#                 doc = fitz.open(file_path)
#                 for page in doc:
#                     pix = page.get_pixmap(dpi=300)
#                     img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
#                     img_byte_arr = io.BytesIO()
#                     img.save(img_byte_arr, format='PNG')
#                     result = reader.readtext(img_byte_arr.getvalue(), detail=0)
#                     text += " ".join(result) + "\n"
#                 doc.close()

#         # PDF
#         elif ext == "pdf":
#             doc = fitz.open(file_path)
#             for page in doc:
#                 text += page.get_text("text") + "\n"
#             doc.close()

#         # DOCX
#         elif ext == "docx":
#             try:
#                 from docx import Document
#                 doc = Document(file_path)
#                 for para in doc.paragraphs:
#                     if para.text.strip():
#                         text += para.text + "\n"
#                 for table in doc.tables:
#                     for row in table.rows:
#                         row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
#                         if row_text:
#                             text += " | ".join(row_text) + "\n"
#             except Exception as e:
#                 print("DOCX extraction error:", e)

#         # TXT
#         elif ext == "txt":
#             with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
#                 text = f.read()

#         else:
#             print(f"Unsupported extension: {ext}, trying PyMuPDF anyway")
#             doc = fitz.open(file_path)
#             for page in doc:
#                 text += page.get_text("text") + "\n"
#             doc.close()

#     except Exception as e:
#         print(f"Extraction error for {ext}: {e}")
#         # Final OCR fallback for any file
#         try:
#             reader = easyocr.Reader(['en', 'hi'], gpu=False)
#             doc = fitz.open(file_path)
#             for page in doc:
#                 pix = page.get_pixmap(dpi=300)
#                 img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
#                 img_byte_arr = io.BytesIO()
#                 img.save(img_byte_arr, format='PNG')
#                 result = reader.readtext(img_byte_arr.getvalue(), detail=0)
#                 text += " ".join(result) + "\n"
#             doc.close()
#         except Exception as ocr_e:
#             print("OCR fallback failed:", ocr_e)

#     final_text = text.strip()
#     print(f"Final extracted text length: {len(final_text)} characters")

#     return final_text[:25000]

def extract_text_from_file(file_path: str) -> str:
    print(f"Extracting from: {file_path}")
    text = ""

    try:
        # PyMuPDF - primary extraction (works well for PDFs and most DOCX)
        import fitz # pyright: ignore[reportMissingImports]
        doc = fitz.open(file_path)
        for page in doc:
            text += page.get_text("text") + "\n"
        doc.close()
    except Exception as e:
        print("PyMuPDF error:", e)

    # Extra DOCX support if needed
    if len(text.strip()) < 300 and file_path.lower().endswith('.docx'):
        try:
            from docx import Document
            doc = Document(file_path)
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n"
            print("DOCX extraction successful")
        except Exception as e:
            print("DOCX error:", e)

    final_text = text.strip()
    print(f"Final extracted length: {len(final_text)} characters")
    return final_text[:25000]

def analyze_contract(text: str, jurisdiction: str = "India") -> dict:
    if not text or len(text) < 100:
        return {
            "overall_risk": 30,
            "summary": "Could not extract readable text from the document.",
            "top_risks": [],
            "clauses": []
        }

    system_prompt = f"""You are ContractBuddy, an expert Indian legal analyst.
Jurisdiction: {jurisdiction}.

Analyze the document and return **ONLY** this exact JSON format. No other text, no markdown, no explanation.

{{
  "overall_risk":  number (0-100),
  "summary": "short summary in 1-2 sentences",
  "top_risks": ["risk 1", "risk 2"],
  "clauses": [
    {{"clause": "clause title", "risk": number, "explanation": "simple explanation"}}
  ]
}}

Document:
{text[:14000]}"""

    try:
        response = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": "Output the JSON now."}
            ],
            temperature=0.1,      # Very low for consistent JSON
            max_tokens=3000
        )

        content = response.choices[0].message.content.strip()

        # Clean common LLM wrappers
        if content.startswith("```json"):
            content = content[7:]
        if content.startswith("```"):
            content = content[3:]
        if content.endswith("```"):
            content = content[:-3]

        content = content.strip()

        result = json.loads(content)
        return result

    except json.JSONDecodeError as e:
        print("JSON Decode Error. Raw output was:", content[:300] if 'content' in locals() else "No content")
    except Exception as e:
        print("Groq error:", str(e))

    # Final safe fallback
    return {
        "overall_risk": 45,
        "summary": "The AI had trouble structuring the analysis. The document appears to be a legal deed. Please try uploading again or use a PDF version if possible.",
        "top_risks": ["Parsing difficulty"],
        "clauses": []
    }